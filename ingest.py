"""
PhotonX RAG - Ingestion Pipeline
Read local documents (.docx) -> Extract -> Chunk -> Embed -> Persist (Chroma)

No web crawling: instead of crawling photonxtech.com, this reads .docx files
placed in SOURCE_DIR (e.g. the PhotonX Company Profile) and indexes them.
Drop additional .docx files into SOURCE_DIR any time and re-run to pick
them up -- it only re-embeds files whose content actually changed
(hash-based diffing), so repeat runs are cheap and safe to put on a cron
job or run after every edit to a source document.

Usage:
    python ingest.py                # incremental update (default)
    python ingest.py --full         # force re-embed everything
    python ingest.py --dir ./docs   # use a different source folder
"""

import argparse
import hashlib
import json
import sys
import time
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)
import chromadb
from sentence_transformers import SentenceTransformer

load_dotenv()

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
SOURCE_DIR = "./source_docs"      # drop .docx files here (e.g. the company profile)

DB_DIR = "./chroma_db"
COLLECTION_NAME = "photonxtech"
METADATA_FILE = "./ingest_metadata.json"

EMBED_MODEL_NAME = "BAAI/bge-base-en-v1.5"
CHUNK_SIZE = 400          # tokens (approx, via char heuristic below)
CHUNK_OVERLAP = 50

HEADERS_TO_SPLIT_ON = [
    ("#", "h1"),
    ("##", "h2"),
    ("###", "h3"),
]

# PhotonX's docx template doesn't necessarily use Word's built-in Heading
# styles for section titles -- section titles like "Who We Are" or "Our
# Services" are often just short, fully-bold paragraphs. We promote those
# to markdown headers too (in addition to real Heading styles) so the
# header-aware splitter below still groups content by section, the same
# way it used to group by the website's h1/h2/h3 markdown headers.
BOLD_HEADING_MAX_CHARS = 80


# ---------------------------------------------------------------------------
# .docx -> structured markdown (order-preserving: paragraphs + tables)
# ---------------------------------------------------------------------------
def _iter_block_items(doc: Document):
    """Yield paragraphs and tables in the order they appear in the document.

    python-docx exposes .paragraphs and .tables as separate flat lists with
    no shared ordering, which would scramble a document that interleaves
    prose and tables (like the "Impact at a Glance" stats table sitting
    between two paragraphs). Walking doc.element.body directly preserves
    the true reading order.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _heading_level(style_name: str) -> int | None:
    if not style_name:
        return None
    style_name = style_name.strip()
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        return min(int(digits), 3) if digits else 1
    return None


def _is_bold_pseudo_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > BOLD_HEADING_MAX_CHARS:
        return False
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)


def _table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    # Rendered as a simple markdown table -- good enough for the LLM to read
    # (stats grids, project tables) without needing a real HTML table.
    lines = ["| " + " | ".join(r) + " |" for r in rows]
    return "\n".join(lines)


def docx_to_markdown(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            table_md = _table_to_markdown(block)
            if table_md:
                lines.append(table_md)
            continue

        paragraph = block
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        level = _heading_level(style_name)

        if level:
            lines.append(f"{'#' * level} {text}")
        elif _is_bold_pseudo_heading(paragraph):
            lines.append(f"## {text}")
        elif text.startswith(("-", "\u2022")) or style_name.lower().startswith("list"):
            clean = text.lstrip("-\u2022* ").strip()
            lines.append(f"- {clean}")
        else:
            lines.append(text)

    return "\n\n".join(lines)


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# ---------------------------------------------------------------------------
# Chunking (same header-aware approach used for the website content before)
# ---------------------------------------------------------------------------
def chunk_markdown(source: str, title: str, markdown: str) -> list[dict]:
    header_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=HEADERS_TO_SPLIT_ON,
        strip_headers=False,
    )
    sections = header_splitter.split_text(markdown)

    fallback_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE * 4,   # rough chars-per-token heuristic
        chunk_overlap=CHUNK_OVERLAP * 4,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []
    for section in sections:
        section_text = section.page_content.strip()
        if not section_text:
            continue

        if len(section_text) <= CHUNK_SIZE * 4:
            pieces = [section_text]
        else:
            pieces = fallback_splitter.split_text(section_text)

        for piece in pieces:
            if len(piece.strip()) < 20:
                continue
            chunks.append(
                {
                    "text": piece.strip(),
                    "source": source,
                    "title": title,
                    "h1": section.metadata.get("h1", ""),
                    "h2": section.metadata.get("h2", ""),
                    "h3": section.metadata.get("h3", ""),
                }
            )

    return chunks


# ---------------------------------------------------------------------------
# Metadata store (for incremental diffing, now keyed by filename)
# ---------------------------------------------------------------------------
def load_metadata() -> dict:
    path = Path(METADATA_FILE)
    if path.exists():
        return json.loads(path.read_text())
    return {}


def save_metadata(metadata: dict) -> None:
    Path(METADATA_FILE).write_text(json.dumps(metadata, indent=2))


# ---------------------------------------------------------------------------
# Main ingestion run
# ---------------------------------------------------------------------------
def run(source_dir: str, full_rebuild: bool = False):
    src_path = Path(source_dir)
    if not src_path.exists():
        print(
            f"Source folder '{source_dir}' does not exist. Create it and drop "
            f".docx files in (e.g. the PhotonX Company Profile), then re-run."
        )
        sys.exit(1)

    docx_files = sorted(src_path.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in '{source_dir}'. Nothing to ingest.")
        sys.exit(1)

    print(f"Found {len(docx_files)} document(s) in {source_dir}.")

    client = chromadb.PersistentClient(path=DB_DIR)

    if full_rebuild:
        # --full previously only reset the in-memory metadata dict, which
        # made every source look "new" to the diffing logic below -- but
        # that logic only deletes a source's old chunks when it sees a
        # *prior* metadata entry for it. With the dict wiped, old chunks
        # never got deleted, so a --full run could silently duplicate
        # content instead of actually starting fresh. Dropping the
        # collection here guarantees --full means "empty slate".
        try:
            client.delete_collection(COLLECTION_NAME)
            print(f"Full rebuild requested: cleared existing '{COLLECTION_NAME}' collection.")
        except Exception:
            pass  # collection didn't exist yet -- nothing to clear

    collection = client.get_or_create_collection(
        COLLECTION_NAME, metadata={"hnsw:space": "cosine"}
    )

    print(f"Loading embedding model: {EMBED_MODEL_NAME} (first run downloads it)")
    embedder = SentenceTransformer(EMBED_MODEL_NAME)

    old_metadata = {} if full_rebuild else load_metadata()
    new_metadata = {}
    seen_sources = set()

    added, updated, unchanged, skipped = 0, 0, 0, 0

    for docx_path in docx_files:
        source = docx_path.name
        seen_sources.add(source)

        try:
            markdown = docx_to_markdown(docx_path)
        except Exception as e:
            print(f"  [warn] Failed to read {source}: {e}")
            skipped += 1
            continue

        if not markdown.strip():
            skipped += 1
            continue

        title = docx_path.stem.replace("_", " ").replace("-", " ").strip()
        h = content_hash(markdown)

        prev = old_metadata.get(source)
        if prev and prev["hash"] == h:
            # Unchanged - carry forward, no re-embedding needed
            new_metadata[source] = prev
            unchanged += 1
            continue

        # Changed or new document: wipe old chunks for this source, re-embed
        if prev:
            collection.delete(where={"source": source})
            updated += 1
        else:
            added += 1

        chunks = chunk_markdown(source, title, markdown)
        if not chunks:
            continue

        texts = [c["text"] for c in chunks]
        embeddings = embedder.encode(texts, normalize_embeddings=True).tolist()
        ids = [f"{source}::chunk::{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "source": c["source"],
                "title": c["title"],
                "h1": c["h1"],
                "h2": c["h2"],
                "h3": c["h3"],
            }
            for c in chunks
        ]

        collection.add(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)

        new_metadata[source] = {
            "hash": h,
            "title": title,
            "chunk_count": len(chunks),
            "last_updated": time.strftime("%Y-%m-%d %H:%M:%S"),
        }
        print(f"  [{'updated' if prev else 'new'}] {source} -> {len(chunks)} chunks")

    # Handle documents removed from the source folder since last run
    removed_sources = set(old_metadata.keys()) - seen_sources
    for source in removed_sources:
        collection.delete(where={"source": source})
        print(f"  [removed] {source}")

    save_metadata(new_metadata)

    print("\n--- Ingestion summary ---")
    print(f"New documents:       {added}")
    print(f"Updated documents:   {updated}")
    print(f"Unchanged documents: {unchanged}")
    print(f"Skipped (empty/failed): {skipped}")
    print(f"Removed documents:   {len(removed_sources)}")
    print(f"Total chunks in collection: {collection.count()}")
    print(f"Persisted to: {DB_DIR}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--full", action="store_true", help="Force full re-embed of all documents"
    )
    parser.add_argument(
        "--dir", default=SOURCE_DIR, help="Folder containing .docx source files"
    )
    args = parser.parse_args()

    run(source_dir=args.dir, full_rebuild=args.full)
